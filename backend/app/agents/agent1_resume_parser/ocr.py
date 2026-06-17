import fitz  # PyMuPDF

from docx import Document

from PIL import (
    Image,
    ImageFilter,
    ImageEnhance,
)

import pytesseract


class OCRProcessor:
    """
    Handles text extraction from:
    - PDF
    - DOCX
    - Images
    """

    @staticmethod
    def extract_text_from_pdf(
        file_path: str,
    ) -> str:

        text_blocks = []

        pdf_document = fitz.open(file_path)

        try:
            for page in pdf_document:

                blocks = page.get_text(
                    "blocks"
                )

                blocks.sort(
                    key=lambda block: (
                        block[1],
                        block[0],
                    )
                )

                for block in blocks:

                    content = (
                        block[4]
                        .strip()
                    )

                    if content:
                        text_blocks.append(
                            content
                        )

        finally:
            pdf_document.close()

        return "\n".join(
            text_blocks
        ).strip()

    @staticmethod
    def extract_text_from_docx(
        file_path: str,
    ) -> str:

        document = Document(
            file_path
        )

        content = []

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:
                content.append(text)

        for table in document.tables:

            for row in table.rows:

                row_data = []

                for cell in row.cells:

                    cell_text = (
                        cell.text.strip()
                    )

                    if cell_text:
                        row_data.append(
                            cell_text
                        )

                if row_data:
                    content.append(
                        " | ".join(
                            row_data
                        )
                    )

        return "\n".join(
            content
        ).strip()

    @staticmethod
    def extract_text_from_image(
        file_path: str,
    ) -> str:

        image = (
            Image.open(file_path)
            .convert("L")
        )

        if image.width < 1800:

            scale = (
                1800
                / image.width
            )

            image = image.resize(
                (
                    int(
                        image.width
                        * scale
                    ),
                    int(
                        image.height
                        * scale
                    ),
                ),
                Image.LANCZOS,
            )

        image = image.filter(
            ImageFilter.SHARPEN
        )

        image = ImageEnhance.Contrast(
            image
        ).enhance(2.0)

        extracted_text = (
            pytesseract
            .image_to_string(
                image,
                config="--oem 3 --psm 6",
            )
        )

        return extracted_text.strip()


def extract_text_from_pdf(
    file_path: str,
) -> str:
    return OCRProcessor.extract_text_from_pdf(
        file_path
    )


def extract_text_from_docx(
    file_path: str,
) -> str:
    return OCRProcessor.extract_text_from_docx(
        file_path
    )


def extract_text_from_image(
    file_path: str,
) -> str:
    return OCRProcessor.extract_text_from_image(
        file_path
    )